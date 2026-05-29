#!/usr/bin/env python3
"""Lista de Compras Inteligente v3 — pure tkinter, dark theme"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import anthropic, base64, json, os, re, shutil
from datetime import datetime
from PIL import Image, ImageTk
import threading
from pathlib import Path

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors as RC
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

try:
    from docx import Document as DocxDoc
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ── Caminhos ──────────────────────────────────────────────────────────────────
BASE_DIR = Path(r"D:\1 - Jefferson\12 - Claude code\Estoque")
IMG_DIR  = BASE_DIR / "images"
DATA_DIR = BASE_DIR / "data"
EXP_DIR  = BASE_DIR / "exports"
CFG_FILE = BASE_DIR / "config.json"
for _d in [IMG_DIR, DATA_DIR, EXP_DIR]:
    _d.mkdir(parents=True, exist_ok=True)

# ── Paleta de cores (Catppuccin Mocha) ────────────────────────────────────────
BG      = "#1e1e2e"
BG2     = "#181825"
BG3     = "#313244"
BG4     = "#45475a"
TEXT    = "#cdd6f4"
SUBT    = "#a6adc8"
GRAY    = "#6c7086"
BORDER  = "#585b70"
BLUE    = "#89b4fa"
BLUE2   = "#6499f5"
GREEN   = "#a6e3a1"
GREEN2  = "#7fc97a"
RED     = "#f38ba8"
RED2    = "#d4728f"
ORANGE  = "#fab387"
ORANGE2 = "#de9a75"
CYAN    = "#89dceb"

# ── Fontes ────────────────────────────────────────────────────────────────────
FS   = ("Segoe UI", 9)
FN   = ("Segoe UI", 10)
FNB  = ("Segoe UI", 10, "bold")
FM   = ("Segoe UI", 11)
FMB  = ("Segoe UI", 11, "bold")
FH   = ("Segoe UI", 13, "bold")
FT   = ("Segoe UI", 17, "bold")

CATEGORIES = [
    "Alimentos","Bebidas","Limpeza","Higiene","Hortifruti",
    "Carnes e Frios","Laticínios","Padaria e Cereais",
    "Congelados","Petshop","Outros",
]
CAT_ICONS = {
    "Alimentos":"🍚","Bebidas":"🥤","Limpeza":"🧹","Higiene":"🪥",
    "Hortifruti":"🥦","Carnes e Frios":"🥩","Laticínios":"🥛",
    "Padaria e Cereais":"🍞","Congelados":"🧊","Petshop":"🐾","Outros":"📦",
}

def cfg_load():
    if CFG_FILE.exists():
        with open(CFG_FILE, encoding="utf-8") as f:
            return json.load(f)
    return {}

def cfg_save(data):
    with open(CFG_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

# ── Botão estilizado ──────────────────────────────────────────────────────────
def mkbtn(parent, text, command, bg=BG3, fg=TEXT, hbg=BG4, width=None, padx=10, pady=6):
    b = tk.Button(parent, text=text, command=command,
                  bg=bg, fg=fg, activebackground=hbg, activeforeground=fg,
                  relief="flat", bd=0, font=FN, cursor="hand2",
                  padx=padx, pady=pady)
    if width:
        b.configure(width=width)
    return b

# ── Frame com rolagem ─────────────────────────────────────────────────────────
class ScrollableFrame(tk.Frame):
    def __init__(self, parent, **kw):
        super().__init__(parent, bg=BG, **kw)
        self._cv = tk.Canvas(self, bg=BG, highlightthickness=0, bd=0)
        self._sb = ttk.Scrollbar(self, orient="vertical", command=self._cv.yview)
        self._cv.configure(yscrollcommand=self._sb.set)
        self.inner = tk.Frame(self._cv, bg=BG)
        self._wid  = self._cv.create_window((0,0), window=self.inner, anchor="nw")
        self._cv.pack(side="left", fill="both", expand=True)
        self._sb.pack(side="right", fill="y")
        self.inner.bind("<Configure>", lambda _:
            self._cv.configure(scrollregion=self._cv.bbox("all")))
        self._cv.bind("<Configure>", lambda e:
            self._cv.itemconfig(self._wid, width=e.width))
        self._cv.bind("<Enter>", lambda _:
            self._cv.bind_all("<MouseWheel>", self._wheel))
        self._cv.bind("<Leave>", lambda _:
            self._cv.unbind_all("<MouseWheel>"))

    def _wheel(self, e):
        self._cv.yview_scroll(-1*(e.delta//120), "units")

    def clear(self):
        for w in self.inner.winfo_children():
            w.destroy()

# ── Separador horizontal colorido ─────────────────────────────────────────────
def hsep(parent, color=BORDER, pady=8):
    tk.Frame(parent, bg=color, height=1).pack(fill="x", pady=pady)

# ── Título de seção ───────────────────────────────────────────────────────────
def section_title(parent, text, color=BLUE, pady=(0,8)):
    row = tk.Frame(parent, bg=BG)
    row.pack(fill="x", pady=pady)
    tk.Frame(row, bg=color, width=4).pack(side="left", fill="y", padx=(0,8))
    tk.Label(row, text=text, bg=BG, fg=color, font=FH).pack(side="left")
    return row

# ── Card (caixa com fundo e título) ──────────────────────────────────────────
def card(parent, title, color=ORANGE, pady=(0,8)):
    outer = tk.Frame(parent, bg=BG3, padx=12, pady=10)
    outer.pack(fill="x", pady=pady)
    hdr = tk.Frame(outer, bg=BG3)
    hdr.pack(fill="x", pady=(0,8))
    tk.Frame(hdr, bg=color, width=3).pack(side="left", fill="y", padx=(0,6))
    tk.Label(hdr, text=title, bg=BG3, fg=color, font=FNB).pack(side="left")
    content = tk.Frame(outer, bg=BG3)
    content.pack(fill="x")
    return outer, content

# ── Entry com borda colorida ──────────────────────────────────────────────────
def mkentry(parent, width=None, font=FM, border_color=BORDER, **kw):
    wrap = tk.Frame(parent, bg=border_color, padx=1, pady=1)
    e = tk.Entry(wrap, bg=BG, fg=TEXT, insertbackground=TEXT,
                 relief="flat", font=font, bd=4, **kw)
    if width:
        e.configure(width=width)
    e.pack()
    return wrap, e

# ── Diálogo de chave API ──────────────────────────────────────────────────────
class ApiKeyDialog(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("Configurar Chave API")
        self.geometry("520x250")
        self.resizable(False, False)
        self.configure(bg=BG)
        self.transient(parent)
        self.grab_set()
        self.result = None
        self._build()
        self.lift(); self.focus_force()

    def _build(self):
        f = tk.Frame(self, bg=BG, padx=32, pady=26)
        f.pack(fill="both", expand=True)
        tk.Label(f, text="🔑  Chave API Anthropic",
                 bg=BG, fg=BLUE, font=FT).pack(pady=(0,6))
        tk.Label(f,
                 text="Necessária para análise de imagens com Claude IA.\n"
                      "Crie sua chave em: console.anthropic.com",
                 bg=BG, fg=SUBT, font=FN, justify="center").pack(pady=(0,14))
        _, self._ent = mkentry(f, width=52, border_color=BLUE)
        self._ent.master.pack(pady=(0,16))
        self._ent.configure(show="*")
        self._ent.focus_set()
        brow = tk.Frame(f, bg=BG)
        brow.pack()
        mkbtn(brow,"💾  Salvar", self._save,
              bg=GREEN, fg=BG2, hbg=GREEN2, width=14).pack(side="left", padx=8)
        mkbtn(brow,"Cancelar", self.destroy,
              bg=BG3, fg=TEXT, hbg=BG4, width=12).pack(side="left", padx=8)
        self._ent.bind("<Return>", lambda _: self._save())

    def _save(self):
        key = self._ent.get().strip()
        if not key:
            messagebox.showwarning("Campo vazio","Digite a chave API.",parent=self)
            return
        if not key.startswith("sk-"):
            messagebox.showwarning("Inválida",
                "A chave deve começar com 'sk-...'",parent=self)
            return
        self.result = key
        c = cfg_load(); c["api_key"] = key; cfg_save(c)
        self.destroy()

# ── Aplicativo principal ──────────────────────────────────────────────────────
class App(tk.Tk):
    _PH = "Nome do produto"

    def __init__(self):
        super().__init__()
        self.title("🛒  Lista de Compras Inteligente")
        self.geometry("1260x880")
        self.minsize(1020, 680)
        self.configure(bg=BG)

        # ttk styles (apenas scrollbar e combobox)
        st = ttk.Style(); st.theme_use("clam")
        st.configure("Vertical.TScrollbar",
                     background=BG3, troughcolor=BG2, arrowcolor=GRAY, bordercolor=BG2)
        st.configure("TCombobox", fieldbackground=BG3, background=BG3,
                     foreground=TEXT, selectbackground=BLUE, selectforeground=BG2,
                     arrowcolor=BLUE)
        st.map("TCombobox", fieldbackground=[("readonly",BG3)],
               foreground=[("readonly",TEXT)])
        st.configure("TCheckbutton", background=BG, foreground=TEXT)
        st.map("TCheckbutton", background=[("active",BG)], foreground=[("active",TEXT)])
        st.configure("TProgressbar", troughcolor=BG3, background=BLUE)

        self.cfg      = cfg_load()
        self.img_path = None
        self.img_tk   = None
        self.detected = []
        self.cart     = []

        self._build_ui()
        self._load_session()
        self.protocol("WM_DELETE_WINDOW", self._on_close)

    # ─── UI ───────────────────────────────────────────────────────────────────

    def _build_ui(self):
        # Cabeçalho
        hdr = tk.Frame(self, bg=BG2, pady=14)
        hdr.pack(fill="x")
        tk.Label(hdr, text="🛒  Lista de Compras Inteligente",
                 bg=BG2, fg=BLUE, font=FT).pack(side="left", padx=20)
        mkbtn(hdr,"⚙  Config API",  lambda: self._cfg_open(),
              bg=BG3, fg=SUBT, hbg=BG4).pack(side="right", padx=6)
        mkbtn(hdr,"💾  Salvar Sessão", lambda: self._save_session(),
              bg=BG3, fg=SUBT, hbg=BG4).pack(side="right")

        tk.Frame(self, bg=BORDER, height=1).pack(fill="x")

        # Corpo
        body = tk.Frame(self, bg=BG, padx=14, pady=12)
        body.pack(fill="both", expand=True)

        left = tk.Frame(body, bg=BG, width=415)
        left.pack(side="left", fill="both")
        left.pack_propagate(False)
        self._build_left(left)

        tk.Frame(body, bg=BORDER, width=1).pack(side="left", fill="y", padx=14)

        right = tk.Frame(body, bg=BG)
        right.pack(side="left", fill="both", expand=True)
        self._build_right(right)

    # ─── Painel esquerdo ──────────────────────────────────────────────────────

    def _build_left(self, p):
        section_title(p, "📸  ANÁLISE DE IMAGEM", BLUE)
        hsep(p, pady=(0,8))

        # Preview
        prev = tk.Frame(p, bg=BORDER)
        prev.pack(fill="x")
        self._lbl_img = tk.Label(prev,
            text="📂\n\nNenhuma imagem selecionada\nClique em 'Selecionar Imagem'",
            bg=BG3, fg=GRAY, font=FN, justify="center", height=12, anchor="center")
        self._lbl_img.pack(padx=1, pady=1, fill="x")

        # Botões
        brow = tk.Frame(p, bg=BG)
        brow.pack(fill="x", pady=8)
        brow.columnconfigure((0,1), weight=1, uniform="a")

        mkbtn(brow,"📂  Selecionar Imagem", lambda: self._pick_image(),
              bg=BLUE, fg=BG2, hbg=BLUE2, pady=7
              ).grid(row=0, column=0, sticky="ew", padx=(0,5))

        self._btn_an = mkbtn(brow,"🔍  Analisar com IA", lambda: self._start_analysis(),
                              bg=BG3, fg=GRAY, hbg=BG4, pady=7)
        self._btn_an.grid(row=0, column=1, sticky="ew", padx=(5,0))
        self._btn_an.configure(state="disabled")

        self._var_st = tk.StringVar(value="Selecione uma imagem para começar")
        tk.Label(p, textvariable=self._var_st, bg=BG, fg=GRAY, font=FS).pack(anchor="w")

        self._pb = ttk.Progressbar(p, mode="indeterminate", length=400)

        hsep(p, pady=10)
        section_title(p, "📋  ITENS DETECTADOS", GREEN, pady=(0,6))

        self._btn_all = mkbtn(p,"✚  Adicionar Todos à Lista",
                               lambda: self._add_all_detected(),
                               bg=GREEN, fg=BG2, hbg=GREEN2, pady=5)
        self._btn_all.pack(fill="x", pady=(0,6))
        self._btn_all.configure(state="disabled")

        hsep(p, pady=(0,6))

        self._sf_det = ScrollableFrame(p)
        self._sf_det.pack(fill="both", expand=True)
        tk.Label(self._sf_det.inner,
                 text="Analise uma imagem para ver\nos itens identificados pela IA",
                 bg=BG, fg=GRAY, font=FN, justify="center").pack(pady=24)

    # ─── Painel direito ───────────────────────────────────────────────────────

    def _build_right(self, p):
        section_title(p, "🛍  LISTA DE COMPRAS", CYAN)
        hsep(p, pady=(0,8))

        # Card de adição manual
        _, cf = card(p,"➕  Adicionar produto manualmente", ORANGE)

        cf.columnconfigure(0, weight=1)
        name_wrap, self._ent_name = mkentry(cf, border_color=BORDER)
        name_wrap.grid(row=0, column=0, sticky="ew", padx=(0,6))
        self._ent_name.insert(0, self._PH)
        self._ent_name.configure(fg=GRAY)
        self._ent_name.bind("<FocusIn>",  self._ph_in)
        self._ent_name.bind("<FocusOut>", self._ph_out)
        self._ent_name.bind("<Return>",   lambda _: self._manual_add())

        self._var_cat = tk.StringVar(value="Alimentos")
        ttk.Combobox(cf, values=CATEGORIES, textvariable=self._var_cat,
                     width=17, state="readonly", font=FN
                     ).grid(row=0, column=1, padx=(0,6), ipady=3)

        qty_wrap, self._ent_qty = mkentry(cf, width=4, border_color=BORDER)
        self._ent_qty.insert(0,"1")
        qty_wrap.grid(row=0, column=2, padx=(0,6))
        self._ent_qty.bind("<Return>", lambda _: self._manual_add())

        mkbtn(cf,"  Adicionar  ", lambda: self._manual_add(),
              bg=ORANGE, fg=BG2, hbg=ORANGE2, pady=5
              ).grid(row=0, column=3)

        # Contador
        self._var_cnt = tk.StringVar(value="0 itens na lista")
        tk.Label(p, textvariable=self._var_cnt,
                 bg=BG, fg=GRAY, font=FS).pack(anchor="e", pady=(0,2))

        # Cabeçalho das colunas
        ch = tk.Frame(p, bg=BG2, pady=5)
        ch.pack(fill="x")
        for txt, wd, expand in [
            ("✓",3,False),("#",3,False),("Produto",0,True),
            ("Categoria",18,False),("Qtd",5,False),("",5,False)
        ]:
            tk.Label(ch, text=txt, width=wd or None,
                     bg=BG2, fg=SUBT, font=FNB, anchor="w"
                     ).pack(side="left", padx=4,
                            **({"expand":True,"fill":"x"} if expand else {}))

        # Scroll da lista
        self._sf_cart = ScrollableFrame(p)
        self._sf_cart.pack(fill="both", expand=True, pady=(0,8))
        self._empty_msg()

        # Exportar
        hsep(p, pady=(0,6))
        exp = tk.Frame(p, bg=BG)
        exp.pack(fill="x")
        tk.Label(exp, text="📤  Exportar:",
                 bg=BG, fg=SUBT, font=FNB).pack(side="left", padx=(0,10))
        for txt, cmd, bg, hbg in [
            ("📄  PDF",   self._export_pdf, RED,    RED2),
            ("📝  TXT",   self._export_txt, GREEN,  GREEN2),
            ("📋  DOC",   self._export_doc, BLUE,   BLUE2),
            ("🗑  Limpar",self._clear_cart, BG3,    BG4),
        ]:
            mkbtn(exp, txt, cmd, bg=bg,
                  fg=BG2 if bg != BG3 else TEXT, hbg=hbg, pady=5
                  ).pack(side="left", padx=4)

    # ─── Placeholder ──────────────────────────────────────────────────────────

    def _ph_in(self, _=None):
        if self._ent_name.get() == self._PH:
            self._ent_name.delete(0,"end")
            self._ent_name.configure(fg=TEXT)

    def _ph_out(self, _=None):
        if not self._ent_name.get().strip():
            self._ent_name.insert(0, self._PH)
            self._ent_name.configure(fg=GRAY)

    def _get_name(self):
        v = self._ent_name.get()
        return "" if v == self._PH else v.strip()

    # ─── Config ───────────────────────────────────────────────────────────────

    def _cfg_open(self):
        dlg = ApiKeyDialog(self)
        self.wait_window(dlg)
        if dlg.result:
            self.cfg["api_key"] = dlg.result

    # ─── Imagem ───────────────────────────────────────────────────────────────

    def _pick_image(self):
        path = filedialog.askopenfilename(
            title="Selecionar imagem",
            filetypes=[("Imagens","*.jpg *.jpeg *.png *.bmp *.webp *.gif"),
                       ("Todos","*.*")])
        if not path:
            return
        dest = IMG_DIR / f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{Path(path).name}"
        shutil.copy2(path, dest)
        self.img_path = str(dest)
        self._preview(path)
        self._btn_an.configure(state="normal", bg=ORANGE, fg=BG2,
                                activebackground=ORANGE2)
        self._var_st.set(f"✓  {Path(path).name}")

    def _preview(self, path):
        try:
            img = Image.open(path)
            img.thumbnail((408, 250), Image.Resampling.LANCZOS)
            self.img_tk = ImageTk.PhotoImage(img)
            self._lbl_img.configure(image=self.img_tk, text="", height=0)
        except Exception as e:
            self._lbl_img.configure(text=f"Erro: {e}", image="")

    # ─── Análise IA ───────────────────────────────────────────────────────────

    def _start_analysis(self):
        api_key = self.cfg.get("api_key")
        if not api_key:
            self._cfg_open()
            api_key = self.cfg.get("api_key")
            if not api_key:
                return
        self._btn_an.configure(state="disabled", text="⏳  Analisando…",
                                bg=BG3, fg=GRAY, activebackground=BG4)
        self._var_st.set("Enviando imagem para Claude IA…")
        self._pb.pack(fill="x", pady=4)
        self._pb.start(12)
        threading.Thread(target=self._do_analysis, args=(api_key,),
                         daemon=True).start()

    def _do_analysis(self, api_key):
        try:
            with open(self.img_path,"rb") as fh:
                b64 = base64.standard_b64encode(fh.read()).decode()
            ext  = Path(self.img_path).suffix.lower()
            mime = {".jpg":"image/jpeg",".jpeg":"image/jpeg",
                    ".png":"image/png",".gif":"image/gif",
                    ".webp":"image/webp"}.get(ext,"image/jpeg")
            prompt = (
                "Você é especialista em análise de estoque. Identifique TODOS os produtos "
                "visíveis na imagem.\n\nResponda SOMENTE com JSON:\n"
                '{"items":[{"nome":"Nome em português","categoria":"Alimentos|Bebidas|'
                'Limpeza|Higiene|Hortifruti|Carnes e Frios|Laticínios|Padaria e Cereais|'
                'Congelados|Petshop|Outros","quantidade_visivel":"estimativa",'
                '"observacao":"opcional"}],"resumo":"resumo geral"}'
            )
            client = anthropic.Anthropic(api_key=api_key)
            msg = client.messages.create(
                model="claude-sonnet-4-6", max_tokens=2048,
                messages=[{"role":"user","content":[
                    {"type":"image","source":{"type":"base64","media_type":mime,"data":b64}},
                    {"type":"text","text":prompt},
                ]}])
            raw = msg.content[0].text.strip()
            m   = re.search(r"\{.*\}", raw, re.DOTALL)
            result = json.loads(m.group() if m else raw)
            self.after(0, self._show_detected, result)
        except Exception as e:
            self.after(0, self._analysis_err, str(e))

    def _show_detected(self, result):
        self._pb.stop(); self._pb.pack_forget()
        self.detected = result.get("items",[])
        resumo        = result.get("resumo","")
        self._sf_det.clear()

        if not self.detected:
            tk.Label(self._sf_det.inner, text="⚠  Nenhum item identificado.",
                     bg=BG, fg=RED, font=FN).pack(pady=20)
        else:
            if resumo:
                tk.Label(self._sf_det.inner, text=resumo,
                         bg=BG, fg=GRAY, font=FS, wraplength=350,
                         justify="left").pack(anchor="w", padx=8, pady=(4,8))
            for item in self.detected:
                self._det_row(item)
            self._btn_all.configure(state="normal")

        self._btn_an.configure(state="normal", text="🔍  Analisar com IA",
                                bg=ORANGE, fg=BG2, activebackground=ORANGE2)
        self._var_st.set(f"✓  {len(self.detected)} item(ns) identificado(s)")

    def _det_row(self, item):
        icon = CAT_ICONS.get(item.get("categoria","Outros"),"📦")
        row  = tk.Frame(self._sf_det.inner, bg=BG3, pady=6, padx=8)
        row.pack(fill="x", pady=2, padx=2)
        info = tk.Frame(row, bg=BG3)
        info.pack(side="left", fill="both", expand=True)
        tk.Label(info, text=f"{icon}  {item['nome']}",
                 bg=BG3, fg=TEXT, font=FMB, anchor="w").pack(anchor="w")
        det = item.get("categoria","")
        if item.get("quantidade_visivel"): det += f"  •  {item['quantidade_visivel']}"
        if item.get("observacao"):         det += f"  •  {item['observacao']}"
        tk.Label(info, text=det, bg=BG3, fg=GRAY, font=FS, anchor="w").pack(anchor="w")
        mkbtn(row, "+ Add",
              lambda it=item: self._cart_add(it["nome"], it.get("categoria","Outros")),
              bg=GREEN, fg=BG2, hbg=GREEN2, width=8, pady=4
              ).pack(side="right")

    def _analysis_err(self, msg):
        self._pb.stop(); self._pb.pack_forget()
        self._btn_an.configure(state="normal", text="🔍  Analisar com IA",
                                bg=ORANGE, fg=BG2, activebackground=ORANGE2)
        self._var_st.set("❌  Erro na análise")
        messagebox.showerror("Erro na Análise",
                             f"Não foi possível analisar.\n\n{msg}\n\nVerifique: ⚙ Config API")

    # ─── Lista de compras ─────────────────────────────────────────────────────

    def _add_all_detected(self):
        for it in self.detected:
            self._cart_add(it["nome"], it.get("categoria","Outros"))

    def _cart_add(self, name, category="Outros", qty=1):
        name = name.strip()
        if not name:
            return
        for it in self.cart:
            if it["nome"].lower() == name.lower():
                it["quantidade"] += qty
                self._refresh_cart()
                return
        self.cart.append({"nome":name,"categoria":category,
                           "quantidade":qty,"checked":False})
        self._refresh_cart()

    def _manual_add(self):
        name = self._get_name()
        if not name:
            messagebox.showwarning("Campo vazio","Digite o nome do produto.")
            return
        try:
            qty = max(1, int(self._ent_qty.get().strip() or "1"))
        except ValueError:
            qty = 1
        self._cart_add(name, self._var_cat.get(), qty)
        self._ent_name.delete(0,"end")
        self._ent_name.insert(0, self._PH)
        self._ent_name.configure(fg=GRAY)
        self._ent_qty.delete(0,"end")
        self._ent_qty.insert(0,"1")
        self._ent_name.focus()

    def _refresh_cart(self):
        self._sf_cart.clear()
        if not self.cart:
            self._empty_msg()
            self._var_cnt.set("0 itens na lista")
            return
        by_cat = {}
        for it in self.cart:
            by_cat.setdefault(it["categoria"],[]).append(it)
        n = 1
        for cat in sorted(by_cat):
            icon = CAT_ICONS.get(cat,"📦")
            ch = tk.Frame(self._sf_cart.inner, bg=BG2, pady=4)
            ch.pack(fill="x", pady=(6,1), padx=2)
            tk.Frame(ch, bg=BLUE, width=3).pack(side="left", fill="y", padx=(4,6))
            tk.Label(ch, text=f"{icon}  {cat}",
                     bg=BG2, fg=BLUE, font=FNB).pack(side="left")
            for it in by_cat[cat]:
                self._cart_row(it, n); n += 1
        self._var_cnt.set(f"{len(self.cart)} item(ns) na lista")

    def _empty_msg(self):
        tk.Label(self._sf_cart.inner,
                 text="Sua lista de compras está vazia.\n\n"
                      "Analise uma imagem ou adicione\nprodutos manualmente.",
                 bg=BG, fg=GRAY, font=FN, justify="center").pack(pady=36)

    def _cart_row(self, item, n):
        row = tk.Frame(self._sf_cart.inner, bg=BG, pady=3)
        row.pack(fill="x", pady=1, padx=2)
        var = tk.BooleanVar(value=item.get("checked",False))
        ttk.Checkbutton(row, variable=var,
                        command=lambda v=var, it=item: self._toggle(it,v)
                        ).pack(side="left", padx=(2,4))
        tk.Label(row, text=str(n), width=3,
                 bg=BG, fg=GRAY, font=FN).pack(side="left")
        tk.Label(row, text=item["nome"],
                 bg=BG, fg=GRAY if item.get("checked") else TEXT,
                 font=FM, width=26, anchor="w").pack(side="left", padx=4)
        tk.Label(row, text=item["categoria"],
                 bg=BG, fg=GRAY, font=FS, width=18, anchor="w").pack(side="left", padx=2)
        qty_var = tk.StringVar(value=str(item["quantidade"]))
        qw, qe  = mkentry(row, width=4, font=FN)
        qe.configure(textvariable=qty_var)
        qw.pack(side="left", padx=4)
        qe.bind("<FocusOut>", lambda _, it=item, v=qty_var: self._set_qty(it,v))
        qe.bind("<Return>",   lambda _, it=item, v=qty_var: self._set_qty(it,v))
        mkbtn(row," ✕ ", lambda it=item: self._remove(it),
              bg=RED, fg=BG2, hbg=RED2, padx=6, pady=3
              ).pack(side="left", padx=2)

    def _toggle(self, item, var):
        item["checked"] = var.get()
        self._refresh_cart()

    def _set_qty(self, item, var):
        try:
            q = int(var.get())
            if q > 0: item["quantidade"] = q
        except ValueError:
            pass

    def _remove(self, item):
        self.cart.remove(item)
        self._refresh_cart()

    def _clear_cart(self):
        if self.cart and messagebox.askyesno("Limpar","Apagar todos os itens?"):
            self.cart.clear(); self._refresh_cart()

    # ─── Exportar ─────────────────────────────────────────────────────────────

    def _warn_empty(self):
        if not self.cart:
            messagebox.showwarning("Lista vazia","Adicione itens antes de exportar.")
            return True
        return False

    def _grouped(self):
        d = {}
        for it in self.cart:
            d.setdefault(it["categoria"],[]).append(it)
        return d

    def _export_txt(self):
        if self._warn_empty(): return
        fname = f"lista_compras_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        sp = filedialog.asksaveasfilename(initialdir=str(EXP_DIR), initialfile=fname,
            defaultextension=".txt", filetypes=[("Texto","*.txt")])
        if not sp: return
        lines = ["="*56,"            LISTA DE COMPRAS",
                 f"   Gerada em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}",
                 "="*56,""]
        for cat, items in sorted(self._grouped().items()):
            lines += [f"  {CAT_ICONS.get(cat,'')}  {cat.upper()}","  "+"─"*44]
            for it in items:
                lines.append(f"  {'[x]' if it.get('checked') else '[ ]'}  "
                              f"{it['nome']:<38}  Qtd: {it['quantidade']}")
            lines.append("")
        lines += ["─"*56, f"  Total: {len(self.cart)} item(ns)", "="*56]
        with open(sp,"w",encoding="utf-8") as fh:
            fh.write("\n".join(lines))
        messagebox.showinfo("Exportado!",f"TXT salvo em:\n{sp}")
        os.startfile(sp)

    def _export_pdf(self):
        if self._warn_empty(): return
        if not REPORTLAB_OK:
            messagebox.showerror("Erro","Execute: pip install reportlab"); return
        fname = f"lista_compras_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        sp = filedialog.asksaveasfilename(initialdir=str(EXP_DIR), initialfile=fname,
            defaultextension=".pdf", filetypes=[("PDF","*.pdf")])
        if not sp: return
        try:
            doc = SimpleDocTemplate(sp, pagesize=A4,
                leftMargin=2*cm, rightMargin=2*cm, topMargin=2.2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet(); story = []
            story.append(Paragraph("Lista de Compras",
                ParagraphStyle("T",parent=styles["Title"],fontSize=22,
                               textColor=RC.HexColor("#1a5276"),spaceAfter=4)))
            story.append(Paragraph(
                f"Gerada em: {datetime.now().strftime('%d/%m/%Y as %H:%M')}",
                ParagraphStyle("D",parent=styles["Normal"],fontSize=10,
                               textColor=RC.grey,spaceAfter=14)))
            hs = ParagraphStyle("H",parent=styles["Heading2"],fontSize=13,
                                textColor=RC.HexColor("#2471a3"),spaceBefore=14,spaceAfter=5)
            for cat, items in sorted(self._grouped().items()):
                story.append(Paragraph(cat, hs))
                data = [["","Produto","Qtd"]]
                for it in items:
                    data.append(["X" if it.get("checked") else "o",
                                  it["nome"], str(it["quantidade"])])
                tbl = Table(data, colWidths=[0.8*cm,13*cm,2.5*cm])
                tbl.setStyle(TableStyle([
                    ("BACKGROUND",(0,0),(-1,0),RC.HexColor("#2471a3")),
                    ("TEXTCOLOR",(0,0),(-1,0),RC.white),
                    ("FONTSIZE",(0,0),(-1,-1),11),
                    ("ROWBACKGROUNDS",(0,1),(-1,-1),[RC.white,RC.HexColor("#eaf4fc")]),
                    ("GRID",(0,0),(-1,-1),0.4,RC.HexColor("#bdc3c7")),
                    ("ALIGN",(0,0),(0,-1),"CENTER"),("ALIGN",(2,0),(2,-1),"CENTER"),
                    ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
                    ("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6),
                    ("LEFTPADDING",(0,0),(-1,-1),8),
                ]))
                story.append(tbl)
            story.append(Spacer(1,0.8*cm))
            story.append(Paragraph(f"Total: {len(self.cart)} item(ns)",
                ParagraphStyle("F",parent=styles["Normal"],fontSize=10,
                               textColor=RC.grey,alignment=1)))
            doc.build(story)
            messagebox.showinfo("Exportado!",f"PDF salvo em:\n{sp}")
            os.startfile(sp)
        except Exception as e:
            messagebox.showerror("Erro ao gerar PDF",str(e))

    def _export_doc(self):
        if self._warn_empty(): return
        if not DOCX_OK:
            messagebox.showerror("Erro","Execute: pip install python-docx"); return
        fname = f"lista_compras_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        sp = filedialog.asksaveasfilename(initialdir=str(EXP_DIR), initialfile=fname,
            defaultextension=".docx", filetypes=[("Word","*.docx")])
        if not sp: return
        try:
            doc = DocxDoc()
            t = doc.add_heading("Lista de Compras",0)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            s = doc.add_paragraph(f"Gerada em: {datetime.now().strftime('%d/%m/%Y as %H:%M')}")
            s.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")
            for cat, items in sorted(self._grouped().items()):
                doc.add_heading(cat, level=2)
                tbl = doc.add_table(rows=1, cols=3)
                tbl.style = "Table Grid"
                hdr = tbl.rows[0].cells
                hdr[0].text, hdr[1].text, hdr[2].text = "","Produto","Qtd"
                for cell in hdr:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                for it in items:
                    r = tbl.add_row().cells
                    r[0].text = "X" if it.get("checked") else "o"
                    r[1].text = it["nome"]; r[2].text = str(it["quantidade"])
                doc.add_paragraph("")
            p = doc.add_paragraph(f"Total: {len(self.cart)} item(ns)")
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            doc.save(sp)
            messagebox.showinfo("Exportado!",f"DOC salvo em:\n{sp}")
            os.startfile(sp)
        except Exception as e:
            messagebox.showerror("Erro ao gerar DOC",str(e))

    # ─── Sessão ───────────────────────────────────────────────────────────────

    def _save_session(self):
        with open(DATA_DIR/"session.json","w",encoding="utf-8") as fh:
            json.dump({"cart":self.cart,"saved":datetime.now().isoformat()},
                      fh, indent=2, ensure_ascii=False)
        self._var_st.set("✓  Sessão salva!")

    def _load_session(self):
        sf = DATA_DIR/"session.json"
        if not sf.exists(): return
        try:
            with open(sf, encoding="utf-8") as fh:
                data = json.load(fh)
            self.cart = data.get("cart",[])
            if self.cart:
                self._refresh_cart()
                saved = data.get("saved","")
                if saved:
                    dt = datetime.fromisoformat(saved)
                    self._var_st.set(f"Sessão carregada ({dt.strftime('%d/%m %H:%M')})")
        except Exception:
            pass

    def _on_close(self):
        if self.cart: self._save_session()
        self.destroy()


if __name__ == "__main__":
    App().mainloop()
